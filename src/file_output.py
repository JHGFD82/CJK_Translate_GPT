"""
File output utilities for the CJK Translation script.
"""

import logging
from pathlib import Path
from typing import Optional
from datetime import datetime

from .config import PDF_MARGINS


def generate_output_filename(input_file: str, source_lang: str, target_lang: str, extension: str = '.txt') -> str:
    """Generate an output filename based on input file and languages."""
    input_path = Path(input_file)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"{input_path.stem}_{source_lang}to{target_lang}_{timestamp}{extension}"
    return str(input_path.parent / output_name)


class FileOutputHandler:
    """Handles saving translations to various file formats."""
    
    @staticmethod
    def save_to_text_file(content: str, output_path: str) -> None:
        """Save content to a text file."""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            logging.info(f'Translation saved to text file: {output_path}')
            print(f"\nTranslation saved to: {output_path}")
        except Exception as e:
            logging.error(f'Error saving to text file: {e}')
            print(f"Error saving to text file: {e}")

    @staticmethod
    def append_to_text_file(content: str, output_path: str) -> None:
        """Append content to a text file."""
        try:
            with open(output_path, 'a', encoding='utf-8') as f:
                f.write(content + '\n\n')
            logging.info(f'Translation appended to text file: {output_path}')
            print(f"Page appended to: {output_path}")
        except Exception as e:
            logging.error(f'Error appending to text file: {e}')
            print(f"Error appending to text file: {e}")
    
    @staticmethod
    def save_to_pdf(content: str, output_path: str, custom_font: Optional[str] = None, target_lang: Optional[str] = None) -> None:
        """Save content to a PDF file using reportlab."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Flowable
            
            # Create PDF document
            doc = SimpleDocTemplate(
                output_path, 
                pagesize=letter,
                rightMargin=PDF_MARGINS['right'],
                leftMargin=PDF_MARGINS['left'],
                topMargin=PDF_MARGINS['top'],
                bottomMargin=PDF_MARGINS['bottom']
            )
            
            # Create story (content container)
            story: list[Flowable] = []
            styles = getSampleStyleSheet()
            
            # Configure font - use Times-Roman for English unless custom font specified
            if not custom_font and target_lang == 'English':
                font_name = 'Times-Roman'
                logging.info(f"Using Times-Roman for English translation (target_lang={target_lang})")
            else:
                font_name = FileOutputHandler._get_cjk_font(custom_font)
                logging.info(f"Using CJK font: {font_name} (custom_font={custom_font}, target_lang={target_lang})")
            
            # Create paragraph style with proper font
            try:
                normal_style = ParagraphStyle(
                    'CJKNormal',
                    parent=styles['Normal'],
                    fontName=font_name,
                    fontSize=12,
                    leading=18,  # 1.5 leading (12pt * 1.5 = 18pt)
                    spaceAfter=12,
                    encoding='utf-8'
                )
                logging.info(f"Created paragraph style with font: {font_name}")
            except Exception as e:
                logging.warning(f"Failed to create custom style with font {font_name}: {e}")
                normal_style = styles['Normal']
                font_name = 'Times-Roman'  # Fallback to default
            
            # Split content into paragraphs and add to story
            paragraphs = content.split('\n\n')
            for i, para in enumerate(paragraphs):
                if para.strip():
                    clean_text = para.strip()
                    try:
                        # For CJK characters, we need to handle the text more carefully
                        # Replace line breaks within paragraphs with spaces
                        clean_text = clean_text.replace('\n', ' ')
                        
                        # Test if the font can handle the text
                        p = Paragraph(clean_text, normal_style)
                        story.append(p)
                        story.append(Spacer(1, 12))
                        logging.debug(f"Successfully added paragraph {i+1} with font {font_name}")
                        
                    except Exception as e:
                        logging.warning(f"Error processing paragraph {i+1} with font {font_name}: {e}")
                        # Try with a different font as fallback
                        try:
                            # Try with Times-Roman as fallback
                            fallback_style = ParagraphStyle(
                                'FallbackCJK',
                                parent=styles['Normal'],
                                fontName='Times-Roman',
                                fontSize=12,
                                leading=18,
                                spaceAfter=12
                            )
                            p = Paragraph(clean_text, fallback_style)
                            story.append(p)
                            story.append(Spacer(1, 12))
                            logging.info(f"Used fallback font Times-Roman for paragraph {i+1}")
                        except Exception as e2:
                            logging.warning(f"Fallback font also failed for paragraph {i+1}: {e2}")
                            # Ultimate fallback: convert to ASCII-safe text
                            ascii_safe_text = clean_text.encode('ascii', 'ignore').decode('ascii')
                            if ascii_safe_text.strip():
                                p = Paragraph(ascii_safe_text, styles['Normal'])
                                story.append(p)
                                story.append(Spacer(1, 12))
                                logging.warning(f"Used ASCII-safe fallback for paragraph {i+1}")
                            else:
                                logging.warning(f"Paragraph {i+1} contained no ASCII-safe characters, skipping")
            
            # Build PDF
            if story:
                doc.build(story)
                logging.info(f'Translation saved to PDF file: {output_path}')
                print(f"\nTranslation saved to PDF: {output_path}")
                if font_name != 'Times-Roman':
                    print(f"Used font: {font_name}")
            else:
                logging.error("No content could be processed for PDF generation")
                print("Error: No content could be processed for PDF generation")
                # Fallback to text file
                text_output_path = output_path.replace('.pdf', '.txt')
                FileOutputHandler.save_to_text_file(content, text_output_path)
            
        except ImportError:
            logging.warning('reportlab not installed. Falling back to text file.')
            print("Warning: reportlab not installed. Saving as text file instead.")
            text_output_path = output_path.replace('.pdf', '.txt')
            FileOutputHandler.save_to_text_file(content, text_output_path)
        except Exception as e:
            logging.error(f'Error saving to PDF: {e}')
            print(f"Error generating PDF: {e}")
            print("Falling back to text file for reliable CJK character support...")
            text_output_path = output_path.replace('.pdf', '.txt')
            FileOutputHandler.save_to_text_file(content, text_output_path)
    
    @staticmethod
    def save_to_docx(content: str, output_path: str, custom_font: Optional[str] = None, target_lang: Optional[str] = None) -> None:
        """Save content to a Word document using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            
            # Create a new document
            doc = Document()
            
            # Set up margins (similar to PDF margins)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
            
            # Get font name for CJK support
            if not custom_font and target_lang == 'English':
                font_name = 'Times New Roman'
                logging.info(f"Using Times New Roman for English translation (target_lang={target_lang})")
            else:
                font_name = FileOutputHandler._get_docx_font(custom_font)
                logging.info(f"Using CJK font for Word: {font_name} (custom_font={custom_font}, target_lang={target_lang})")
            
            # Split content into paragraphs and add to document
            paragraphs = content.split('\n\n')
            for i, para in enumerate(paragraphs):
                if para.strip():
                    clean_text = para.strip()
                    # Replace line breaks within paragraphs with spaces
                    clean_text = clean_text.replace('\n', ' ')
                    
                    try:
                        # Add paragraph to document
                        p = doc.add_paragraph(clean_text)
                        
                        # Configure paragraph formatting
                        paragraph_format = p.paragraph_format
                        paragraph_format.space_after = Pt(12)
                        paragraph_format.line_spacing = 1.5
                        
                        # Configure font
                        for run in p.runs:
                            run.font.name = font_name
                            run.font.size = Pt(12)
                        
                        # If paragraph has no runs (empty), add one with the font
                        if not p.runs:
                            run = p.runs[0] if p.runs else p.add_run()
                            run.font.name = font_name
                            run.font.size = Pt(12)
                        
                        logging.debug(f"Successfully added paragraph {i+1} with font {font_name}")
                        
                    except Exception as e:
                        logging.warning(f"Error processing paragraph {i+1} for Word document: {e}")
                        # Try adding paragraph without special formatting
                        try:
                            p = doc.add_paragraph(clean_text)
                            logging.info(f"Added paragraph {i+1} with basic formatting")
                        except Exception as e2:
                            logging.warning(f"Failed to add paragraph {i+1} to Word document: {e2}")
                            continue
            
            # Save the document
            if len(doc.paragraphs) > 0:
                doc.save(output_path)
                logging.info(f'Translation saved to Word document: {output_path}')
                print(f"\nTranslation saved to Word document: {output_path}")
                if font_name != 'Times New Roman':
                    print(f"Used font: {font_name}")
            else:
                logging.error("No content could be processed for Word document generation")
                print("Error: No content could be processed for Word document generation")
                # Fallback to text file
                text_output_path = output_path.replace('.docx', '.txt')
                FileOutputHandler.save_to_text_file(content, text_output_path)
            
        except ImportError:
            logging.warning('python-docx not installed. Falling back to text file.')
            print("Warning: python-docx not installed. To enable Word document export, install it with:")
            print("pip install python-docx")
            print("Saving as text file instead.")
            text_output_path = output_path.replace('.docx', '.txt')
            FileOutputHandler.save_to_text_file(content, text_output_path)
        except Exception as e:
            logging.error(f'Error saving to Word document: {e}')
            print(f"Error generating Word document: {e}")
            print("Falling back to text file for reliable CJK character support...")
            text_output_path = output_path.replace('.docx', '.txt')
            FileOutputHandler.save_to_text_file(content, text_output_path)
    
    @staticmethod
    def save_translation_output(content: str, input_file: Optional[str], output_file: Optional[str], 
                              auto_save: bool, source_lang: str, target_lang: str, custom_font: Optional[str] = None) -> None:
        """Save translation output to file based on user preferences."""
        if not content.strip():
            print("No content to save.")
            return
        
        # Determine output file path
        if output_file:
            output_path = output_file
        elif auto_save and input_file:
            output_path = generate_output_filename(input_file, source_lang, target_lang, '.txt')
        else:
            # No saving requested
            return
        
        # Ensure directory exists
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Determine file type and save accordingly
        if output_path.lower().endswith('.pdf'):
            FileOutputHandler.save_to_pdf(content, output_path, custom_font, target_lang)
        elif output_path.lower().endswith('.docx'):
            FileOutputHandler.save_to_docx(content, output_path, custom_font, target_lang)
        else:
            # Default to text file
            if not output_path.lower().endswith('.txt'):
                output_path += '.txt'
            FileOutputHandler.save_to_text_file(content, output_path)

    @staticmethod
    def save_page_progressively(content: str, input_file: Optional[str], output_file: Optional[str], 
                               auto_save: bool, source_lang: str, target_lang: str, is_first_page: bool = False,
                               custom_font: Optional[str] = None) -> Optional[str]:
        """Save a single page progressively to output file. Returns the output path."""
        if not content.strip():
            print("No content to save.")
            return None
        
        # Determine output file path
        if output_file:
            output_path = output_file
        elif auto_save and input_file:
            output_path = generate_output_filename(input_file, source_lang, target_lang, '.txt')
        else:
            # No saving requested
            return None
        
        # Ensure directory exists
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # For progressive saving, we currently only support text files
        # PDF and Word document merging is complex and requires additional dependencies
        if output_path.lower().endswith('.pdf'):
            print("Note: Progressive saving for PDF format not yet supported. Using text format.")
            output_path = output_path.replace('.pdf', '.txt')
        elif output_path.lower().endswith('.docx'):
            print("Note: Progressive saving for Word document format not yet supported. Using text format.")
            output_path = output_path.replace('.docx', '.txt')
        
        # Default to text file
        if not output_path.lower().endswith('.txt'):
            output_path += '.txt'
        
        # Save first page or append subsequent pages
        if is_first_page:
            FileOutputHandler.save_to_text_file(content, output_path)
        else:
            FileOutputHandler.append_to_text_file(content, output_path)
        
        return output_path
    
    @staticmethod
    def _get_cjk_font(custom_font: Optional[str] = None) -> str:
        """Get an appropriate font for CJK characters from the fonts directory."""
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import os
            
            # Check fonts from the local fonts directory
            fonts_dir = os.path.join(os.path.dirname(__file__), '..', 'fonts')
            if os.path.exists(fonts_dir):
                # If a custom font is specified, try it first
                if custom_font:
                    custom_font_path = os.path.join(fonts_dir, f"{custom_font}.ttf")
                    if os.path.exists(custom_font_path):
                        try:
                            custom_font_name = f"CustomFont_{custom_font}"
                            if custom_font_name not in pdfmetrics.getRegisteredFontNames():
                                pdfmetrics.registerFont(TTFont(custom_font_name, custom_font_path))  # type: ignore
                            logging.info(f"Using custom CJK font: {custom_font_name}")
                            return custom_font_name
                        except Exception as e:
                            logging.warning(f"Failed to register custom font {custom_font}: {e}")
                            print(f"Warning: Custom font '{custom_font}' failed to load. Using default font selection.")
                    else:
                        logging.warning(f"Custom font file not found: {custom_font_path}")
                        print(f"Warning: Custom font '{custom_font}.ttf' not found in fonts/ directory. Using default font selection.")
                
                # Preferred CJK fonts (in order of preference) - map to available fonts
                preferred_fonts = [
                    ('Arial Unicode.ttf', 'ArialUnicode'),         # Your available Arial Unicode
                    ('AppleGothic.ttf', 'AppleGothic'),           # Apple Gothic (good for CJK)
                    ('AppleMyungjo.ttf', 'AppleMyungjo'),         # Apple Myungjo (good for CJK)
                ]
                
                # First, try preferred fonts
                for font_filename, font_name in preferred_fonts:
                    font_path = os.path.join(fonts_dir, font_filename)
                    if os.path.exists(font_path):
                        try:
                            if font_name not in pdfmetrics.getRegisteredFontNames():
                                pdfmetrics.registerFont(TTFont(font_name, font_path))  # type: ignore
                            logging.info(f"Using preferred CJK font: {font_name} ({font_filename})")
                            return font_name
                        except Exception as e:
                            logging.warning(f"Failed to register preferred font {font_name}: {e}")
                            continue
                
                # If no preferred fonts, use any available .ttf file (reportlab only supports TTF, not OTF)
                for font_file in os.listdir(fonts_dir):
                    if font_file.endswith('.ttf'):
                        font_path = os.path.join(fonts_dir, font_file)
                        # Create a safe font name by removing problematic characters
                        safe_font_name = font_file.replace('.ttf', '').replace('-', '_').replace(',', '_').replace(' ', '_')
                        try:
                            if safe_font_name not in pdfmetrics.getRegisteredFontNames():
                                pdfmetrics.registerFont(TTFont(safe_font_name, font_path))  # type: ignore
                            logging.info(f"Using available CJK font: {safe_font_name} ({font_file})")
                            return safe_font_name
                        except Exception as e:
                            logging.warning(f"Failed to register font {safe_font_name}: {e}")
                            continue
            
            # No compatible fonts found - give clear guidance
            logging.warning("No CJK fonts found in fonts/ directory.")
            print("Warning: No CJK fonts available for PDF generation.")
            print("To fix: Add CJK .ttf fonts to the 'fonts/' directory in this project.")
            print("Note: Only .ttf fonts are supported. OTF fonts will not work with reportlab.")
            print("Recommended CJK fonts:")
            print("  - Arial Unicode MS (Microsoft)")
            print("  - Source Han Sans (Adobe): https://github.com/adobe-fonts/source-han-sans")
            print("  - Apple system fonts (AppleGothic, AppleMyungjo)")
            print("Alternative: Save as .txt file for proper CJK character display.")
            
            return 'Times-Roman'  # Fallback to reportlab default (Times New Roman equivalent)
            
        except Exception as e:
            logging.warning(f"Error checking CJK fonts: {e}")
            return 'Times-Roman'

    @staticmethod
    def _get_docx_font(custom_font: Optional[str] = None) -> str:
        """Get an appropriate font for CJK characters for Word documents."""
        try:
            import os
            
            # Check fonts from the local fonts directory
            fonts_dir = os.path.join(os.path.dirname(__file__), '..', 'fonts')
            if os.path.exists(fonts_dir):
                # If a custom font is specified, try it first
                if custom_font:
                    custom_font_path = os.path.join(fonts_dir, f"{custom_font}.ttf")
                    if os.path.exists(custom_font_path):
                        # For Word documents, we can use the font name directly
                        # Since Word handles font loading differently than reportlab
                        logging.info(f"Using custom CJK font for Word: {custom_font}")
                        return custom_font
                    else:
                        logging.warning(f"Custom font file not found: {custom_font_path}")
                        print(f"Warning: Custom font '{custom_font}.ttf' not found in fonts/ directory. Using default font selection.")
                
                # Preferred CJK fonts for Word documents (system font names)
                preferred_word_fonts = [
                    'Arial Unicode MS',      # Microsoft's Unicode font
                    'AppleGothic',          # Apple Gothic (good for CJK)
                    'AppleMyungjo',         # Apple Myungjo (good for CJK)
                    'Arial',                # Fallback to Arial
                    'Calibri',              # Default Word font
                ]
                
                # For Word documents, we use system font names rather than file paths
                # Word will handle font resolution internally
                for font_name in preferred_word_fonts:
                    logging.info(f"Using CJK font for Word: {font_name}")
                    return font_name
            
            # Default fallback
            logging.info("Using Times New Roman as fallback for Word document")
            return 'Times New Roman'
            
        except Exception as e:
            logging.warning(f"Error checking fonts for Word document: {e}")
            return 'Times New Roman'
